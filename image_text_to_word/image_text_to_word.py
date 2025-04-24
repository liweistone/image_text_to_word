import os
from docx import Document
from docx.shared import Inches
from PIL import Image

def create_word_from_images_and_texts(folder_path, output_docx="output.docx"):
    doc = Document()
    files = os.listdir(folder_path)
    
    # 分离并排序文件（按数字部分）
    image_files = [f for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    text_files = [f for f in files if f.lower().endswith('.txt')]
    
    image_files.sort(key=lambda x: int(''.join(filter(str.isdigit, x))))
    text_files.sort(key=lambda x: int(''.join(filter(str.isdigit, x))))
    
    # 检查数量
    if len(image_files) != len(text_files):
        print(f"警告：图片({len(image_files)})和文本({len(text_files)})数量不匹配！")
    
    # 处理文件
    for i, (img_file, txt_file) in enumerate(zip(image_files, text_files), 1):
        if i % 10 == 0:
            print(f"处理进度：{i}/{len(image_files)}")
        
        # 插入图片
        img_path = os.path.join(folder_path, img_file)
        try:
            img = Image.open(img_path)
            doc.add_picture(img_path, width=Inches(6.0))  # 限制宽度为 6 英寸
        except Exception as e:
            print(f"图片 {img_file} 插入失败: {e}")
            continue
        
        # 插入文本
        txt_path = os.path.join(folder_path, txt_file)
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()
            doc.add_paragraph(text)
        except Exception as e:
            print(f"文本 {txt_file} 读取失败: {e}")
            continue
        
        doc.add_paragraph()  # 空行分隔
    
    doc.save(output_docx)
    print(f"Word 文档已生成：{output_docx}")

if __name__ == "__main__":
    folder_path = input("input_files").strip()
    output_name = input("out").strip() or "output.docx"
    create_word_from_images_and_texts(folder_path, output_name)