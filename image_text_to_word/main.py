import os
from docx import Document
from docx.shared import Inches
from PIL import Image
from config import INPUT_DIR, OUTPUT_DIR

def process_files():
    """主处理逻辑"""
    doc = Document()
    files = os.listdir(INPUT_DIR)
    
    # 分离并排序文件
    image_files = sorted(
        [f for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))],
        key=lambda x: int(''.join(filter(str.isdigit, x)))
    )
    text_files = sorted(
        [f for f in files if f.lower().endswith('.txt')],
        key=lambda x: int(''.join(filter(str.isdigit, x)))
    )

    # 检查文件匹配
    if len(image_files) != len(text_files):
        print(f"⚠️ 警告：图片({len(image_files)})和文本({len(text_files)})数量不匹配")

    # 处理文件
    for idx, (img_file, txt_file) in enumerate(zip(image_files, text_files), 1):
        print(f"\r🚀 处理进度: {idx}/{len(image_files)}", end="", flush=True)
        
        # 插入图片
        img_path = os.path.join(INPUT_DIR, img_file)
        try:
            with Image.open(img_path) as img:
                doc.add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            print(f"\n❌ 图片 {img_file} 插入失败: {e}")
            continue

        # 插入文本
        txt_path = os.path.join(INPUT_DIR, txt_file)
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                doc.add_paragraph(f.read())
        except Exception as e:
            print(f"\n❌ 文本 {txt_file} 读取失败: {e}")
            continue

        doc.add_paragraph()  # 分隔空行

    # 保存文档
    output_path = os.path.join(OUTPUT_DIR, "output.docx")
    doc.save(output_path)
    print(f"\n✅ 文档生成成功！保存至: {output_path}")

if __name__ == "__main__":
    print("="*50)
    print(f"📂 输入目录: {INPUT_DIR}")
    print(f"💾 输出目录: {OUTPUT_DIR}")
    print("="*50)
    
    if not os.listdir(INPUT_DIR):
        print("⚠️ 输入目录为空！请将图片和文本放入 input_files 文件夹")
    else:
        process_files()